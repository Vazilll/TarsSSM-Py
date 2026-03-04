"""
═══════════════════════════════════════════════════════════════
  TARS Document Tools — PDF / Word / Excel Processing
═══════════════════════════════════════════════════════════════

Позволяет ТАРС:
  - Читать PDF, DOCX, XLSX файлы → извлекать текст
  - Создавать и дополнять Word документы
  - Создавать Excel таблицы
  - Загружать документы в LEANN память

Зависимости (устанавливаются лениво):
  pip install PyPDF2 python-docx openpyxl

Usage:
  from tools.document_tools import DocumentReadTool, DocumentWriteTool, SpreadsheetTool
  
  registry.register(DocumentReadTool())
  registry.register(DocumentWriteTool())
  registry.register(SpreadsheetTool())
"""

import os
import json
import csv
import io
import logging
import traceback
from pathlib import Path
from typing import Dict, Any, List, Optional
from concurrent.futures import ThreadPoolExecutor

from tools import Tool, ToolResult

logger = logging.getLogger("Tars.DocumentTools")


# ═══════════════════════════════════════════════════════════════
# Module cache — import ОДИН раз, потом из кеша (0ns)
# ═══════════════════════════════════════════════════════════════

_CACHE: Dict[str, Any] = {}  # module_name → module object


def _ensure_pypdf2():
    if 'PyPDF2' not in _CACHE:
        try:
            import PyPDF2
            _CACHE['PyPDF2'] = PyPDF2
        except ImportError:
            raise ImportError("PyPDF2 not installed. Run: pip install PyPDF2")
    return _CACHE['PyPDF2']


def _ensure_docx():
    if 'docx' not in _CACHE:
        try:
            import docx
            _CACHE['docx'] = docx
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
    return _CACHE['docx']


def _ensure_openpyxl():
    if 'openpyxl' not in _CACHE:
        try:
            import openpyxl
            _CACHE['openpyxl'] = openpyxl
        except ImportError:
            raise ImportError("openpyxl not installed. Run: pip install openpyxl")
    return _CACHE['openpyxl']


# Thread pool для параллельного I/O (PDF pages, multiple files)
_IO_POOL = ThreadPoolExecutor(max_workers=4, thread_name_prefix='tars_doc')


# ═══════════════════════════════════════════════════════════════
# 1. Document Reader — PDF, DOCX, XLSX, TXT
# ═══════════════════════════════════════════════════════════════

def read_pdf(file_path: str, max_pages: int = 50) -> str:
    """
    Извлечь текст из PDF.
    Оптимизация: параллельная экстракция страниц через ThreadPool.
    """
    PyPDF2 = _ensure_pypdf2()
    
    with open(file_path, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        n_pages = min(len(reader.pages), max_pages)
        
        if n_pages == 0:
            return "[PDF пуст]"
        
        # Маленькие PDF — последовательно (нет overhead на пул)
        if n_pages <= 3:
            parts = []
            for i in range(n_pages):
                page_text = reader.pages[i].extract_text() or ""
                if page_text.strip():
                    parts.append(f"--- Страница {i+1} ---\n{page_text}")
            return "\n\n".join(parts) if parts else "[PDF: нет текста]"
        
        # Большие PDF — параллельно по страницам
        def _extract_page(page_idx):
            try:
                text = reader.pages[page_idx].extract_text() or ""
                return (page_idx, text.strip())
            except Exception:
                return (page_idx, "")
        
        futures = [_IO_POOL.submit(_extract_page, i) for i in range(n_pages)]
        results = sorted((f.result() for f in futures), key=lambda x: x[0])
        
        parts = []
        for idx, text in results:
            if text:
                parts.append(f"--- Страница {idx+1} ---\n{text}")
    
    return "\n\n".join(parts) if parts else "[PDF: нет извлекаемого текста]"


def read_docx(file_path: str) -> str:
    """
    Извлечь текст из DOCX (Word).
    Оптимизация: pre-allocated list, minimal style checks.
    """
    docx_mod = _ensure_docx()
    
    doc = docx_mod.Document(file_path)
    
    # Pre-allocate estimated capacity
    n_paras = len(doc.paragraphs)
    n_tables = len(doc.tables)
    parts = [None] * (n_paras + n_tables * 5)  # rough estimate
    idx = 0
    
    # Параграфы (inline style check — fastest path)
    for para in doc.paragraphs:
        text = para.text
        if not text or not text.strip():
            continue
        text = text.strip()
        style_name = para.style.name if para.style else ""
        if 'Heading' in style_name or 'Title' in style_name:
            # Extract digit from style name fast
            level = next((int(c) for c in style_name if c.isdigit()), 1)
            parts[idx] = f"{'#' * level} {text}"
        else:
            parts[idx] = text
        idx += 1
    
    # Таблицы — bulk cell extraction
    for i, table in enumerate(doc.tables):
        parts[idx] = f"\n--- Таблица {i+1} ---"
        idx += 1
        for row in table.rows:
            if idx >= len(parts):
                parts.extend([None] * 100)
            parts[idx] = " | ".join(c.text.strip() for c in row.cells)
            idx += 1
    
    # Trim None slots and join
    return "\n".join(p for p in parts[:idx] if p is not None) or "[Документ пуст]"


def read_xlsx(file_path: str, max_rows: int = 500) -> str:
    """
    Извлечь данные из XLSX (Excel).
    Оптимизация: read_only mode, bulk iter_rows, StringIO buffer.
    """
    openpyxl = _ensure_openpyxl()
    
    wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
    buf = io.StringIO()
    
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        buf.write(f"═══ Лист: {sheet_name} ═══\n")
        
        row_count = 0
        header_written = False
        
        for row in ws.iter_rows(values_only=True, max_row=max_rows + 1):
            if row_count > max_rows:
                remaining = (ws.max_row or max_rows) - max_rows
                if remaining > 0:
                    buf.write(f"... ({remaining} строк пропущено)\n")
                break
            
            line = " | ".join(str(c) if c is not None else "" for c in row)
            buf.write(line)
            buf.write('\n')
            
            if row_count == 0 and not header_written:
                # Separator after header
                buf.write("-|-".join("---" for _ in row))
                buf.write('\n')
                header_written = True
            
            row_count += 1
        
        if row_count == 0:
            buf.write("[Лист пуст]\n")
    
    wb.close()
    return buf.getvalue()


def _read_text_fast(file_path: str, max_bytes: int = 10000) -> str:
    """Быстрое чтение текстового файла (mmap для больших файлов)."""
    size = os.path.getsize(file_path)
    
    if size <= max_bytes:
        # Маленький файл — читаем целиком
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            return f.read()
    
    # Большой файл — читаем только max_bytes через буфер
    with open(file_path, 'r', encoding='utf-8', errors='replace',
              buffering=65536) as f:
        return f.read(max_bytes)


def _read_csv_fast(file_path: str, max_rows: int = 500) -> str:
    """Быстрое чтение CSV через буферизованный IO."""
    buf = io.StringIO()
    with open(file_path, 'r', encoding='utf-8', errors='replace',
              buffering=65536) as f:
        reader = csv.reader(f)
        for i, row in enumerate(reader):
            if i >= max_rows:
                buf.write(f"... (ещё строк: {i}+)\n")
                break
            buf.write(" | ".join(row))
            buf.write('\n')
    result = buf.getvalue()
    return result if result else "[CSV пуст]"


def _read_json_fast(file_path: str) -> str:
    """Быстрое чтение JSON с лимитом вывода."""
    with open(file_path, 'r', encoding='utf-8', buffering=65536) as f:
        data = json.load(f)
    return json.dumps(data, indent=2, ensure_ascii=False)[:10000]


# Dispatch table — O(1) lookup вместо if-elif chain
_READER_DISPATCH = {
    '.pdf': lambda fp, mp, mr: read_pdf(fp, mp),
    '.docx': lambda fp, mp, mr: read_docx(fp),
    '.doc': lambda fp, mp, mr: "[.doc формат устарел. Пересохраните как .docx]",
    '.xlsx': lambda fp, mp, mr: read_xlsx(fp, mr),
    '.xls': lambda fp, mp, mr: "[.xls формат устарел. Пересохраните как .xlsx]",
    '.csv': lambda fp, mp, mr: _read_csv_fast(fp, mr),
    '.json': lambda fp, mp, mr: _read_json_fast(fp),
    '.xml': lambda fp, mp, mr: _read_text_fast(fp),
    '.txt': lambda fp, mp, mr: _read_text_fast(fp),
    '.md': lambda fp, mp, mr: _read_text_fast(fp),
    '.rst': lambda fp, mp, mr: _read_text_fast(fp),
    '.log': lambda fp, mp, mr: _read_text_fast(fp),
    '.ini': lambda fp, mp, mr: _read_text_fast(fp),
    '.cfg': lambda fp, mp, mr: _read_text_fast(fp),
    '.yml': lambda fp, mp, mr: _read_text_fast(fp),
    '.yaml': lambda fp, mp, mr: _read_text_fast(fp),
}


def read_document(file_path: str, max_pages: int = 50, max_rows: int = 500) -> str:
    """
    Универсальный читатель документов.
    Dispatch table O(1) для всех форматов.
    """
    path = Path(file_path)
    
    if not path.exists():
        raise FileNotFoundError(f"Файл не найден: {file_path}")
    
    ext = path.suffix.lower()
    handler = _READER_DISPATCH.get(ext)
    
    if handler:
        return handler(file_path, max_pages, max_rows)
    
    # Fallback — try text
    try:
        return _read_text_fast(file_path, 5000)
    except Exception:
        return f"[Неподдерживаемый формат: {ext}]"


# ═══════════════════════════════════════════════════════════════
# 2. Document Writer — DOCX
# ═══════════════════════════════════════════════════════════════

def write_docx(
    file_path: str,
    content: List[Dict[str, Any]],
    title: str = "",
    append: bool = False,
) -> str:
    """
    Создать или дополнить Word документ.
    
    Args:
        file_path: путь к .docx файлу
        content: список блоков, каждый:
            {"type": "heading", "text": "Заголовок", "level": 1}
            {"type": "paragraph", "text": "Текст параграфа"}
            {"type": "bullet", "items": ["Пункт 1", "Пункт 2"]}
            {"type": "table", "headers": ["A", "B"], "rows": [["1", "2"]]}
            {"type": "image", "path": "image.png", "width_cm": 15}
            {"type": "page_break"}
        title: заголовок документа (для новых документов)
        append: если True — дополнить существующий файл
    
    Returns:
        str — путь к созданному файлу
    """
    docx_mod = _ensure_docx()
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Load existing or create new
    if append and os.path.exists(file_path):
        doc = docx_mod.Document(file_path)
    else:
        doc = docx_mod.Document()
        if title:
            doc.add_heading(title, level=0)
    
    for block in content:
        block_type = block.get("type", "paragraph")
        
        if block_type == "heading":
            level = block.get("level", 1)
            doc.add_heading(block.get("text", ""), level=level)
        
        elif block_type == "paragraph":
            text = block.get("text", "")
            p = doc.add_paragraph(text)
            if block.get("bold"):
                for run in p.runs:
                    run.bold = True
            if block.get("italic"):
                for run in p.runs:
                    run.italic = True
            if block.get("align") == "center":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        elif block_type == "bullet":
            items = block.get("items", [])
            for item in items:
                doc.add_paragraph(item, style='List Bullet')
        
        elif block_type == "numbered":
            items = block.get("items", [])
            for item in items:
                doc.add_paragraph(item, style='List Number')
        
        elif block_type == "table":
            headers = block.get("headers", [])
            rows = block.get("rows", [])
            
            if headers:
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Table Grid'
                
                # Header row
                for i, h in enumerate(headers):
                    cell = table.rows[0].cells[i]
                    cell.text = str(h)
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True
                
                # Data rows
                for row_data in rows:
                    row = table.add_row()
                    for i, val in enumerate(row_data):
                        if i < len(headers):
                            row.cells[i].text = str(val)
        
        elif block_type == "image":
            img_path = block.get("path", "")
            width = Cm(block.get("width_cm", 15))
            if os.path.exists(img_path):
                doc.add_picture(img_path, width=width)
            else:
                doc.add_paragraph(f"[Изображение не найдено: {img_path}]")
        
        elif block_type == "page_break":
            doc.add_page_break()
    
    # Ensure directory exists
    os.makedirs(os.path.dirname(file_path) or ".", exist_ok=True)
    doc.save(file_path)
    
    return file_path


# ═══════════════════════════════════════════════════════════════
# 3. Spreadsheet Writer — XLSX
# ═══════════════════════════════════════════════════════════════

def write_xlsx(
    file_path: str,
    sheets: List[Dict[str, Any]],
) -> str:
    """
    Создать Excel файл.
    
    Args:
        file_path: путь к .xlsx файлу
        sheets: список листов:
            {
                "name": "Лист1",
                "headers": ["Колонка A", "Колонка B"],
                "rows": [["значение1", "значение2"], ...],
                "column_widths": [20, 30],  # optional
                "freeze_header": True,       # optional
            }
    
    Returns:
        str — путь к созданному файлу
    """
    openpyxl = _ensure_openpyxl()
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    
    wb = openpyxl.Workbook()
    # Remove default sheet if we're creating named ones
    if sheets:
        wb.remove(wb.active)
    
    # Styles
    header_font = Font(bold=True, size=11, color="FFFFFF")
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin'),
    )
    alt_fill = PatternFill(start_color="D9E2F3", end_color="D9E2F3", fill_type="solid")
    
    for sheet_data in sheets:
        sheet_name = sheet_data.get("name", "Sheet1")
        headers = sheet_data.get("headers", [])
        rows = sheet_data.get("rows", [])
        widths = sheet_data.get("column_widths", [])
        freeze = sheet_data.get("freeze_header", True)
        
        ws = wb.create_sheet(title=sheet_name)
        
        # Header row
        if headers:
            for col_idx, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col_idx, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_align
                cell.border = thin_border
        
        # Data rows
        for row_idx, row_data in enumerate(rows, 2):
            for col_idx, value in enumerate(row_data, 1):
                cell = ws.cell(row=row_idx, column=col_idx, value=_auto_type(value))
                cell.border = thin_border
                cell.alignment = Alignment(wrap_text=True)
                # Alternating row colors
                if row_idx % 2 == 0:
                    cell.fill = alt_fill
        
        # Column widths
        if widths:
            for i, w in enumerate(widths):
                ws.column_dimensions[get_column_letter(i + 1)].width = w
        else:
            # Auto-width (approximate)
            for col_idx in range(1, len(headers) + 1):
                max_len = len(str(headers[col_idx - 1])) if col_idx <= len(headers) else 10
                for row in rows:
                    if col_idx - 1 < len(row):
                        max_len = max(max_len, len(str(row[col_idx - 1])))
                ws.column_dimensions[get_column_letter(col_idx)].width = min(max_len + 4, 50)
        
        # Freeze header row
        if freeze and headers:
            ws.freeze_panes = "A2"
        
        # Auto-filter
        if headers:
            ws.auto_filter.ref = f"A1:{get_column_letter(len(headers))}{len(rows) + 1}"
    
    # Save
    os.makedirs(os.path.dirname(file_path) or ".", exist_ok=True)
    wb.save(file_path)
    
    return file_path


def _auto_type(value):
    """Автоматическое определение типа для Excel ячейки."""
    if value is None or value == "":
        return None
    
    s = str(value).strip()
    
    # Integer
    try:
        if '.' not in s and 'e' not in s.lower():
            return int(s)
    except ValueError:
        pass
    
    # Float
    try:
        return float(s)
    except ValueError:
        pass
    
    # Bool
    if s.lower() in ("true", "да", "yes"):
        return True
    if s.lower() in ("false", "нет", "no"):
        return False
    
    return s


# ═══════════════════════════════════════════════════════════════
# 4. TARS Tool Wrappers (for ToolRegistry)
# ═══════════════════════════════════════════════════════════════

class DocumentReadTool(Tool):
    """
    Инструмент чтения документов.
    Поддерживает: PDF, DOCX, XLSX, CSV, JSON, TXT, MD, XML
    """
    
    def name(self) -> str:
        return "read_document"
    
    def description(self) -> str:
        return (
            "Прочитать содержимое документа. "
            "Поддерживает форматы: PDF, DOCX (Word), XLSX (Excel), "
            "CSV, JSON, TXT, MD, XML. "
            "Извлекает текст, таблицы и структуру документа."
        )
    
    def parameters(self) -> Dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "Абсолютный путь к файлу документа",
                },
                "max_pages": {
                    "type": "integer",
                    "description": "Макс. страниц для PDF (default: 50)",
                },
                "ingest_to_memory": {
                    "type": "boolean",
                    "description": "Загрузить в LEANN память после чтения (default: false)",
                },
            },
            "required": ["file_path"],
        }
    
    async def execute(self, args: Dict[str, Any]) -> ToolResult:
        file_path = args.get("file_path", "")
        max_pages = args.get("max_pages", 50)
        ingest = args.get("ingest_to_memory", False)
        
        if not file_path:
            return ToolResult.error("file_path is required")
        
        try:
            text = read_document(file_path, max_pages=max_pages)
            
            # Truncate for LLM output if too long
            llm_text = text[:4000]
            if len(text) > 4000:
                llm_text += f"\n... [усечено, всего {len(text)} символов]"
            
            # Optionally ingest to LEANN memory
            if ingest:
                try:
                    from memory.leann import LeannIndex
                    leann = LeannIndex()
                    
                    # Chunk long documents
                    chunk_size = 500
                    chunks = []
                    filename = Path(file_path).name
                    for i in range(0, len(text), chunk_size):
                        chunk = text[i:i + chunk_size]
                        if chunk.strip():
                            chunks.append(f"[{filename}] {chunk}")
                    
                    for chunk in chunks:
                        leann.add_document(chunk)
                    
                    leann.save()
                    llm_text += f"\n\n✓ Загружено в память: {len(chunks)} чанков"
                    logger.info(f"Ingested {len(chunks)} chunks from {file_path}")
                except Exception as e:
                    llm_text += f"\n\n⚠ Ошибка загрузки в память: {e}"
            
            return ToolResult.success(
                output=f"Прочитано: {file_path} ({len(text)} символов)",
                llm_output=llm_text,
            )
            
        except ImportError as e:
            return ToolResult.error(str(e))
        except Exception as e:
            logger.error(f"Error reading {file_path}: {e}")
            return ToolResult.error(f"Ошибка чтения: {e}")


class DocumentWriteTool(Tool):
    """
    Инструмент создания/дополнения Word документов (.docx).
    """
    
    def name(self) -> str:
        return "write_document"
    
    def description(self) -> str:
        return (
            "Создать или дополнить Word документ (.docx). "
            "Поддерживает: заголовки, параграфы, списки, таблицы, "
            "изображения, разрывы страниц. "
            "Можно дополнять существующий документ (append=true)."
        )
    
    def parameters(self) -> Dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "Путь к .docx файлу для создания/дополнения",
                },
                "title": {
                    "type": "string",
                    "description": "Заголовок документа (для новых документов)",
                },
                "append": {
                    "type": "boolean",
                    "description": "Дополнить существующий файл (default: false)",
                },
                "content": {
                    "type": "array",
                    "description": (
                        "Массив блоков контента. Каждый блок — объект с полем 'type'. "
                        "Типы: 'heading' (text, level), 'paragraph' (text, bold, italic), "
                        "'bullet' (items[]), 'numbered' (items[]), "
                        "'table' (headers[], rows[][]), "
                        "'image' (path, width_cm), 'page_break'"
                    ),
                    "items": {"type": "object"},
                },
            },
            "required": ["file_path", "content"],
        }
    
    async def execute(self, args: Dict[str, Any]) -> ToolResult:
        file_path = args.get("file_path", "")
        title = args.get("title", "")
        append = args.get("append", False)
        content = args.get("content", [])
        
        if not file_path:
            return ToolResult.error("file_path is required")
        
        if not file_path.endswith('.docx'):
            file_path += '.docx'
        
        if not content:
            return ToolResult.error("content is required (array of blocks)")
        
        try:
            result_path = write_docx(file_path, content, title=title, append=append)
            
            action = "дополнен" if append else "создан"
            return ToolResult.success(
                output=f"Документ {action}: {result_path} ({len(content)} блоков)",
                llm_output=f"Word документ {action}: {result_path}",
            )
            
        except ImportError as e:
            return ToolResult.error(str(e))
        except Exception as e:
            logger.error(f"Error writing docx: {e}\n{traceback.format_exc()}")
            return ToolResult.error(f"Ошибка записи: {e}")


class SpreadsheetTool(Tool):
    """
    Инструмент создания Excel таблиц (.xlsx).
    """
    
    def name(self) -> str:
        return "create_spreadsheet"
    
    def description(self) -> str:
        return (
            "Создать Excel таблицу (.xlsx). "
            "Поддерживает: несколько листов, заголовки, автоширина колонок, "
            "чередование цветов строк, заморозка заголовков, автофильтры. "
            "Данные автоматически типизируются (числа, даты, булевы)."
        )
    
    def parameters(self) -> Dict[str, Any]:
        return {
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "Путь к .xlsx файлу",
                },
                "sheets": {
                    "type": "array",
                    "description": (
                        "Массив листов. Каждый лист: "
                        "{name, headers[], rows[][], column_widths[], freeze_header}"
                    ),
                    "items": {
                        "type": "object",
                        "properties": {
                            "name": {"type": "string"},
                            "headers": {"type": "array", "items": {"type": "string"}},
                            "rows": {"type": "array"},
                            "column_widths": {"type": "array", "items": {"type": "integer"}},
                            "freeze_header": {"type": "boolean"},
                        },
                    },
                },
            },
            "required": ["file_path", "sheets"],
        }
    
    async def execute(self, args: Dict[str, Any]) -> ToolResult:
        file_path = args.get("file_path", "")
        sheets = args.get("sheets", [])
        
        if not file_path:
            return ToolResult.error("file_path is required")
        
        if not file_path.endswith('.xlsx'):
            file_path += '.xlsx'
        
        if not sheets:
            return ToolResult.error("sheets is required (array of sheet data)")
        
        try:
            result_path = write_xlsx(file_path, sheets)
            
            total_rows = sum(len(s.get("rows", [])) for s in sheets)
            return ToolResult.success(
                output=(
                    f"Excel создан: {result_path} "
                    f"({len(sheets)} лист(ов), {total_rows} строк)"
                ),
                llm_output=f"Excel файл создан: {result_path}",
            )
            
        except ImportError as e:
            return ToolResult.error(str(e))
        except Exception as e:
            logger.error(f"Error writing xlsx: {e}\n{traceback.format_exc()}")
            return ToolResult.error(f"Ошибка записи: {e}")
